#!/usr/bin/env python3
import os
import re
import json
import time
import logging
from datetime import date
from datetime import datetime
from io import BytesIO
from typing import List, Dict

import streamlit as st
import google.generativeai as genai
import openai
import pandas as pd
import torch # Added torch import for granular analysis logic
from dotenv import load_dotenv
from transformers import AutoTokenizer, AutoModelForSequenceClassification, pipeline
from PyPDF2 import PdfReader
import docx
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

# =========================
# Config & LLMOps Logging
# =========================
LOG_FILE = "metrics_log.jsonl"
logging.basicConfig(
    filename="smartlegal_metrics.log",
    level=logging.INFO,
    format="%(asctime)s %(levelname)s %(message)s",
)

def _ensure_metrics_state():
    if "metrics" not in st.session_state:
        st.session_state.metrics = {
            "total_requests": 0,
            "failed_requests": 0,
            "total_latency": 0.0,
            "avg_latency": 0.0,
        }

def _bump_success(latency_s: float):
    _ensure_metrics_state()
    st.session_state.metrics["total_requests"] += 1
    st.session_state.metrics["total_latency"] += latency_s
    tr = st.session_state.metrics["total_requests"]
    st.session_state.metrics["avg_latency"] = (
        st.session_state.metrics["total_latency"] / max(tr, 1)
    )

def _bump_failure():
    _ensure_metrics_state()
    st.session_state.metrics["failed_requests"] += 1
    st.session_state.metrics["total_requests"] += 1

def log_metric(event_type, latency, tokens, cost, status="SUCCESS", model="gemini-2.0-flash"):
    entry = {
        "timestamp": datetime.utcnow().isoformat(),
        "event": event_type,
        "model": model,
        "latency_s": latency,
        "tokens_used": tokens,
        "cost_gbp": cost,
        "status": status,
    }
    with open(LOG_FILE, "a") as f:
        f.write(json.dumps(entry) + "\n")
    if status == "SUCCESS":
        _bump_success(latency)
    else:
        _bump_failure()
    logging.info(entry)

# =========================
# Document Formatting Helper
# =========================
def create_formatted_agreement(draft_text, tenant, landlord):
    doc = Document()

    title = doc.add_paragraph("RENTAL AGREEMENT")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.runs[0]
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(16)

    subtitle = doc.add_paragraph("(Under the Model Tenancy Act, 2021)")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(12)
    doc.add_paragraph()

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(12)
    normal_style.paragraph_format.line_spacing = 1.15

    clean_text = re.sub(r"#+", "", draft_text)
    clean_text = re.sub(r"\*\*", "", clean_text)
    lines = [ln.strip() for ln in clean_text.split("\n") if ln.strip()]

    for line in lines:
        heading_match = re.match(r"^(WHEREAS|NOW THEREFORE|IN WITNESS|THIS RENTAL AGREEMENT|BETWEEN|AND|IMPORTANT NOTES|SIGNED)(.*)", line, re.IGNORECASE)
        if heading_match:
            keyword = heading_match.group(1).strip()
            rest = heading_match.group(2).strip()

            p = doc.add_paragraph()
            run_bold = p.add_run(keyword)
            run_bold.bold = True

            if rest:
                p.add_run(" " + rest)

            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)

        elif re.match(r"^[0-9]+\.", line):
            # Match patterns like "1. Rent: The monthly rent ..."
            clause_match = re.match(r"^([0-9]+\.\s*)([^:]+)(:?\s*)(.*)", line)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

            if clause_match:
                num = clause_match.group(1)          # "1. "
                title = clause_match.group(2).strip()  # "Rent"
                after_colon = clause_match.group(3)   # ":" or whitespace
                rest = clause_match.group(4).strip()  # rest of the text

                p.add_run(num)  # number normal
                bold_run = p.add_run(title)
                bold_run.bold = True
                if after_colon:
                    p.add_run(after_colon)
                if rest:
                    p.add_run(rest)
            else:
                # fallback if pattern not matched
                p.add_run(line)

        elif line.startswith("*"):
            clean = line.lstrip("* ").strip()
            p = doc.add_paragraph(clean, style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.6)

        else:
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(6)

    doc.add_page_break()
    doc.add_paragraph(
        "IN WITNESS WHEREOF, the parties hereto have executed this Agreement.",
        style="Normal",
    )
    table = doc.add_table(rows=2, cols=2)
    table.autofit = True

    table.cell(0, 0).text = (
        f"By the Landlord:\n\n(Signature)\n\n{landlord}\n"
        f"DOB: {st.session_state.get('landlord_dob_input', '')}\n"
        f"Address: {st.session_state.get('landlord_address', '')}"
    )
    table.cell(0, 1).text = (
        f"By the Tenant:\n\n(Signature)\n\n{tenant}\n"
        f"DOB: {st.session_state.get('tenant_dob_input', '')}\n"
        f"Address: {st.session_state.get('tenant_address', '')}"
    )

    filename = f"Formatted_Rental_{tenant}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(filename)
    return filename


# =========================
# Splitting the Document for Risk Analysis
# =========================

LEGAL_MODEL_PATH = os.path.join("models", "rental-risk-bert")

def split_text_into_clauses(text: str) -> List[str]:
    """
    Splits the full legal document text into sentences or short clauses,
    filtering out signatures, headers, and witnessing blocks.
    """
    # 1. Split the text into potential clauses using punctuation separators
    # Ensures splitting occurs at sentence boundaries.
    clauses = re.split(r'(?<!\w\.\w.)(?<![A-Z][a-z]\.)(?<=\.|\?|\!|;)\s+(?=[A-Z])', text)
    
    # Define keywords/phrases to filter out: headers, footers, and signature blocks
    filter_patterns = [
        # Headers/Titles/Metadata/Placeholders
        r'^\s*(RENTAL AGREEMENT|TENANCY ACT|DOB|Address|Landlord|Tenant|Date of Birth|\[Agreed Number\]|\[Date\]|\[Agreed Amount\])\s*$', 
        # Witnessing/Execution blocks
        r'^\s*(IN WITNESS WHEREOF|By the Landlord|By the Tenant|Signature|EXECUTED|AGREEMENT PREAMBLE)\s*',
        # Clause numbers/short fragments
        r'^\s*\d+\.\s*$',
        # Empty lines or very short phrases (under 10 characters)
        r'^.{0,10}$', 
        # Section titles that don't end in punctuation
        r'^(SECTION|ARTICLE|PREAMBLE)\s+\d+',
    ]

    # 2. Filter the clauses
    filtered_clauses = []
    
    for clause in clauses:
        clause_clean = clause.strip()
        
        # Skip if the clause is too short or is an empty string
        if len(clause_clean) < 10:
            continue
            
        # Check against all filter patterns
        is_metadata = False
        for pattern in filter_patterns:
            if re.match(pattern, clause_clean, re.IGNORECASE):
                is_metadata = True
                break
        
        if not is_metadata:
            filtered_clauses.append(clause_clean)

    return filtered_clauses

@st.cache_resource
def load_risk_model_for_granular_analysis():
    """Loads the model and tokenizer explicitly for detailed analysis."""
    try:
        # Load the model and tokenizer from the saved directory
        tokenizer = AutoTokenizer.from_pretrained(LEGAL_MODEL_PATH)
        model = AutoModelForSequenceClassification.from_pretrained(LEGAL_MODEL_PATH)
        return tokenizer, model
    except Exception as e:
        st.error(f"Failed to load fine-tuned model: {e}")
        return None, None

def analyze_document_risk_granular(document_text):
    """
    Loads the fine-tuned model and analyzes a document clause-by-clause,
    returning a detailed list of risks.
    """
    tokenizer, model = load_risk_model_for_granular_analysis()
    if not model:
        return {"error": "Risk model not loaded."}

    # Get the label mappings saved during training
    id2label = model.config.id2label
    
    # 1. Split the document into clauses
    clauses = split_text_into_clauses(document_text)
    
    if not clauses:
        return {"error": "Could not extract any recognizable clauses for analysis."}

    # 2. Tokenize and predict
    inputs = tokenizer(
        clauses,
        padding=True,
        truncation=True,
        max_length=256,
        return_tensors="pt"
    )

    with torch.no_grad():
        outputs = model(**inputs)

    # 3. Process results
    logits = outputs.logits
    probabilities = torch.softmax(logits, dim=1).tolist()
    predictions = torch.argmax(logits, dim=1).tolist()
    
    # 4. Compile final analysis
    analysis_results = []
    risky_count = 0

    for clause, pred_id, probs in zip(clauses, predictions, probabilities):
        label = id2label[pred_id]
        confidence = probs[pred_id]
        
        analysis_results.append({
            "clause": clause,
            "prediction": label,
            "confidence": f"{confidence * 100:.2f}%",
            "is_risky": label == "RISKY"
        })
        if label == "RISKY":
            risky_count += 1
    
    return {
        "summary": f"Total Clauses Analyzed: {len(analysis_results)}. Risky Clauses Detected: {risky_count}.",
        "details": analysis_results,
        "risky_count": risky_count
    }


# =========================
# App Init & Models
# =========================
load_dotenv()
genai.configure(api_key=os.getenv("GEMINI_API_KEY"))
openai.api_key = os.getenv("OPENAI_API_KEY")

@st.cache_resource
def load_risk_pipeline():
    """
    Loads the fine-tuned SmartLegal rental risk classification model.
    Uses the locally trained DistilBERT model under models/rental-risk-bert/.
    """
    return pipeline(
        "text-classification",
        model=LEGAL_MODEL_PATH,
        tokenizer=LEGAL_MODEL_PATH,
        truncation=True,
        max_length=512,
    )

# Load the model explicitly for caching purposes, 
load_risk_pipeline() 

st.title("SmartLegal Rental Assistant")
st.markdown("**Draft • Review • Fix — Based on Model Tenancy Act 2021**")

# --- Persistent tab handling ---
tab_names = ["Draft New Agreement", "Review & Fix Agreement"]
selected_tab = st.session_state.get("selected_tab", tab_names[0])
selected_tab = st.radio("Navigation", tab_names, horizontal=True,
                        index=tab_names.index(selected_tab), key="tab_selector")
st.session_state.selected_tab = selected_tab

# --- Persistent file + results state ---
if "uploaded_file" not in st.session_state:
    st.session_state.uploaded_file = None
if "review_results" not in st.session_state:
    st.session_state.review_results = {}

# =========================
# TAB 1: Draft New Agreement
# =========================
if selected_tab == "Draft New Agreement":
    st.header("Create New Rental Agreement")

    with st.form("rental_form", clear_on_submit=False):
        landlord = st.text_input("Landlord Name", key="landlord_name")
        landlord_address = st.text_area("Landlord Address", key="landlord_address")
        landlord_dob = st.date_input("Landlord Date of Birth", key="landlord_dob_input",
                                     min_value=date(1900, 1, 1), max_value=date.today())

        tenant = st.text_input("Tenant Name", key="tenant_name")
        tenant_address = st.text_area("Tenant Address", key="tenant_address")
        tenant_dob = st.date_input("Tenant Date of Birth", key="tenant_dob_input",
                                   min_value=date(1900, 1, 1), max_value=date.today())

        rent = st.number_input("Monthly Rent ₹", min_value=1000, key="rent_amount")
        deposit = st.number_input("Security Deposit ₹", min_value=0, key="deposit_amount")
        address = st.text_area("Property Address", key="property_address")
        start = st.date_input("Start Date", key="lease_start_input",
                              min_value=date.today(), max_value=date(2100, 12, 31))
        months = st.selectbox("Duration", ["11 months", "2 years", "3 years"], key="lease_duration")
        amenities = st.text_area("Amenities (optional)", key="amenities_text")

        submitted = st.form_submit_button("Generate Agreement")

    if submitted:
        if not all([landlord, landlord_address, landlord_dob, tenant, tenant_address, tenant_dob, address]):
            st.error("Please fill in all mandatory fields: Landlord and Tenant names, addresses, dates of birth, and property address.")
        else:
            with st.spinner('Generating rental agreement... Please wait.'):
                prompt = f"""
                Draft a complete rental agreement under the Model Tenancy Act 2021 with the following details:

                Landlord:
                  Name: {landlord}
                  Date of Birth: {landlord_dob}
                  Address: {landlord_address}

                Tenant:
                  Name: {tenant}
                  Date of Birth: {tenant_dob}
                  Address: {tenant_address}

                Property: {address}
                Rent: ₹{rent}/month
                Deposit: ₹{deposit}
                Duration: {months}, starting from {start}.
                Amenities: {amenities if amenities else 'None'}

                Include all mandatory clauses such as registration, police verification,
                maintenance, notice period, and eviction.

                IMPORTANT: Do not include any placeholders, examples, or 'e.g.' text.
                Use the exact values provided above and write in clear, formal legal language.
                """

                try:
                    model = genai.GenerativeModel("models/gemini-2.0-flash")
                    t0 = time.time()
                    response = model.generate_content(prompt)
                    t1 = time.time()
                    draft = (response.text or "").strip()

                    for prefix in ("sure", "okay", "here", "below", "this is"):
                        if draft.lower().startswith(prefix):
                            parts = draft.split("\n", 1)
                            draft = parts[1].strip() if len(parts) > 1 else draft
                            break

                    latency = round(t1 - t0, 2)
                    token_count = len(prompt.split()) + len(draft.split())
                    cost = round(token_count * 0.0005 / 1000, 6)
                    st.caption(f"Latency: {latency}s | Tokens: {token_count} | Cost: £{cost}")
                    log_metric("DraftAgreement", latency, token_count, cost, status="SUCCESS")
                except Exception as e:
                    st.error(f"Gemini API error: {e}")
                    draft = "Unable to generate agreement text."
                    log_metric("DraftAgreement", 0, 0, 0, status="FAILED")

            st.success("Agreement Drafted Successfully!")
            st.text_area("Preview", draft, height=400, key="draft_preview")

            with st.spinner('Creating formatted document...'):
                file_name = create_formatted_agreement(draft, tenant, landlord)

            with open(file_name, "rb") as f:
                st.download_button(label="Download Formatted Word File",
                                   data=f, file_name=file_name,
                                   mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# =========================
# TAB 2: Review & Fix Agreement
# =========================
elif selected_tab == "Review & Fix Agreement":
    st.header("Review & Suggest Amendments")

    uploaded = st.file_uploader(
        "Upload PDF / DOCX / TXT",
        type=["pdf", "docx", "txt"],
        key="review_uploader"
    )

    if "file_cache" not in st.session_state:
        st.session_state.file_cache = {}
    if "current_file_key" not in st.session_state:
        st.session_state.current_file_key = None

    # Helper to build a stable file key for caching
    def make_file_key(name, size):
        return f"{name.replace(' ', '_')}_{size}"

    # If user uploaded a new file right now, capture it into session_state.file_cache
    if uploaded is not None:
        file_key = make_file_key(uploaded.name, uploaded.size)
        st.session_state.current_file_key = file_key

        # If not already cached, read bytes and create cache entry
        if file_key not in st.session_state.file_cache:
            with st.spinner("Saving uploaded file to session cache..."):
                try:
                    raw = uploaded.getvalue()
                except Exception:
                    uploaded.seek(0)
                    raw = uploaded.read()
                st.session_state.file_cache[file_key] = {
                    "name": uploaded.name,
                    "size": uploaded.size,
                    "bytes": raw,
                }

    # If uploader is empty but we have a current_file_key in session, keep using it
    file_key = st.session_state.current_file_key
    file_entry = st.session_state.file_cache.get(file_key) if file_key else None

    if not file_entry:
        st.info("Upload a PDF/DOCX/TXT to review.")
    else:
        if "text" not in st.session_state.file_cache[file_key]:
            with st.spinner("Processing uploaded document..."):
                data = file_entry["bytes"]
                lower = file_entry["name"].lower()
                try:
                    if lower.endswith(".pdf"):
                        reader = PdfReader(BytesIO(data))
                        text = " ".join([(p.extract_text() or "") for p in reader.pages])
                    elif lower.endswith(".docx") or lower.endswith(".doc"):
                        bio = BytesIO(data)
                        d = docx.Document(bio)
                        text = "\n".join([p.text for p in d.paragraphs])
                    else:
                        text = data.decode("utf-8", errors="ignore")
                except Exception as e:
                    st.error(f"Document parsing error: {e}")
                    text = ""
                st.session_state.file_cache[file_key]["text"] = text

        text = st.session_state.file_cache[file_key].get("text", "")

        model = genai.GenerativeModel("models/gemini-2.0-flash")

        # Summary
        if "summary" not in st.session_state.file_cache[file_key]:
            with st.spinner("Generating summary..."):
                t0 = time.time()
                summary = model.generate_content(f"Summarize this rental agreement in 100 words:\n{text[:4000]}").text
                t1 = time.time()
                latency_summary = round(t1 - t0, 2)
                tokens_summary = len(summary.split()) + len(text.split()[:4000])
                cost_summary = round(tokens_summary * 0.0005 / 1000, 6)
                st.session_state.file_cache[file_key]["summary"] = summary
                st.session_state.file_cache[file_key]["summary_metrics"] = (latency_summary, tokens_summary, cost_summary)
                log_metric("Summarisation", latency_summary, tokens_summary, cost_summary, status="SUCCESS")

        summary = st.session_state.file_cache[file_key]["summary"]
        latency_summary, tokens_summary, cost_summary = st.session_state.file_cache[file_key]["summary_metrics"]
        st.subheader("Summary")
        st.text_area("Summary Preview", summary, height=150, key=f"review_summary_{file_key}")
        st.caption(f"Latency: {latency_summary}s | Tokens: {tokens_summary} | Cost: £{cost_summary}")

        # Risk (now on-demand via button)
        run_risk = st.button("Run Risk Assessment")

        if run_risk:
            with st.spinner('Analyzing risk clause-by-clause...'):
                analysis_results = analyze_document_risk_granular(text[:10000])
                
                if "error" in analysis_results:
                    st.error(f"Risk analysis error: {analysis_results['error']}")
                else:
                    st.session_state.file_cache[file_key]["analysis"] = analysis_results
                    log_metric("RiskClassification", 0, 0, 0, status="SUCCESS", model="distilbert-sst2")
            
        if "analysis" in st.session_state.file_cache[file_key]:
            analysis_results = st.session_state.file_cache[file_key]["analysis"]
            risky_count = analysis_results["risky_count"]
            total_clauses = len(analysis_results["details"])
            
            st.subheader("Granular Risk Assessment")
            
            # Display overall risk summary
            if risky_count > 0:
                st.warning(f"**HIGH RISK:** Detected {risky_count} out of {total_clauses} clauses as problematic.")
            else:
                st.success(f"LOW RISK: All {total_clauses} analyzed clauses appear safe.")
            
            # Display detailed clause breakdown
            st.markdown("---")
            st.markdown("##### Clause-by-Clause Details:")
            
            for detail in analysis_results["details"]:
                clause = detail['clause']
                pred = detail['prediction']
                conf = detail['confidence']
                
                if pred == "RISKY":
                    st.error(f"**RISKY ({conf}):** {clause}")
                else:
                    st.markdown(f"**SAFE ({conf}):** *{clause}*")
            
        # Amendments
        if "amendments" not in st.session_state.file_cache[file_key]:
            with st.spinner("Generating amendment suggestions..."):
                t0 = time.time()
                amendments = model.generate_content(
                    f"""List missing or incorrect clauses according to Model Tenancy Act 2021.
                    Do not include any conversational phrases, introductions, or explanations.
                    {text[:5000]}
                    """
                ).text
                t1 = time.time()
                latency_amend = round(t1 - t0, 2)
                tokens_amend = len(amendments.split()) + len(text.split()[:5000])
                cost_amend = round(tokens_amend * 0.0005 / 1000, 6)
                amendments = re.sub(r"^\s*[\*\-•]\s*", "", amendments, flags=re.MULTILINE)
                st.session_state.file_cache[file_key]["amendments"] = (amendments, latency_amend, tokens_amend, cost_amend)
                log_metric("AmendmentReview", latency_amend, tokens_amend, cost_amend, status="SUCCESS")

        amendments, latency_amend, tokens_amend, cost_amend = st.session_state.file_cache[file_key]["amendments"]
        st.subheader("Suggested Amendments")
        st.text_area("Amendments Preview", amendments, height=200, key=f"review_amendments_{file_key}")
        st.caption(f"Latency: {latency_amend}s | Tokens: {tokens_amend} | Cost: £{cost_amend}")

        st.session_state.success_count = st.session_state.get("success_count", 0) + 1
        st.sidebar.metric("Total Successful Queries", st.session_state.success_count)

# =========================
# Sidebar: Voice + Metrics
# =========================
with st.sidebar:
    st.header("Voice to Clause")
    audio = st.file_uploader("Upload voice note", type=["mp3", "wav"])
    if audio:
        try:
            transcript = openai.audio.transcriptions.create(model="whisper-1", file=audio).text
            st.code(transcript)
            clause = openai.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{
                    "role": "user",
                    "content": f"Convert to legal clause (Model Tenancy Act 2021): {transcript}"
                }],
            ).choices[0].message.content
            st.success(clause)
            log_metric("VoiceToClause", 0, len(transcript.split()), 0.0001, status="SUCCESS", model="whisper+gpt-4o-mini")
        except Exception as e:
            st.error(f"Voice-to-clause error: {e}")
            log_metric("VoiceToClause", 0, 0, 0, status="FAILED", model="whisper+gpt-4o-mini")

    st.subheader("📈 LLMOps Metrics Summary")
    _ensure_metrics_state()
    st.metric("Total Requests", st.session_state.metrics["total_requests"])
    st.metric("Avg Latency (s)", round(st.session_state.metrics["avg_latency"], 2))
    st.metric("Failed Requests", st.session_state.metrics["failed_requests"])

    if os.path.exists(LOG_FILE):
        try:
            df = pd.read_json(LOG_FILE, lines=True)
            st.dataframe(df.tail(5), use_container_width=True)
            st.caption(f"Total Logs: {len(df)} | Last Updated: {df.iloc[-1]['timestamp']}")
        except Exception:
            st.info("Metrics log exists but could not be parsed yet.")
    else:
        st.info("No metrics logged yet. Generate an agreement to start.")
